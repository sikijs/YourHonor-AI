"""Document export: convert content (markdown or HTML) to PDF, DOCX, or raw markdown.

- PDF is rendered with fpdf2's HTML renderer. fpdf2 ships only Latin-1
  built-in fonts, so text is normalized to Latin-1-safe characters first
  (curly quotes, em/en dashes, accents, etc.) to keep PDF generation fully
  offline and dependency-free inside the Docker image.
- DOCX is built with python-docx.
- Markdown export returns the content as-is.

The API layer never exposes raw errors here; ExportError maps to HTTP 400.
"""

import re
import unicodedata
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from fpdf import FPDF

SUPPORTED_FORMATS = ("pdf", "docx", "md")
SUPPORTED_CONTENT_TYPES = ("markdown", "html")


class ExportError(ValueError):
    pass


# Tags fpdf2's write_html understands; everything else is unwrapped before
# rendering so stray <div>/<span>/<section> markup (as produced by the
# frontend print helpers) never breaks PDF generation.
_FPDF_SAFE_TAGS = {
    "p", "br", "hr", "h1", "h2", "h3", "h4", "h5", "h6",
    "b", "strong", "i", "em", "u", "a", "ul", "ol", "li",
    "table", "tr", "td", "th", "blockquote", "code", "pre",
}

_SMART_CHAR_MAP = {
    "\u2018": "'", "\u2019": "'",  # left/right single quotation marks
    "\u201c": '"', "\u201d": '"',  # left/right double quotation marks
    "\u2013": "-", "\u2014": "-",  # en/em dashes
    "\u2026": "...",              # ellipsis
}


def _sanitize_latin1(text: str) -> str:
    """Normalize smart punctuation and accents into Latin-1-safe ASCII."""
    text = "".join(_SMART_CHAR_MAP.get(ch, ch) for ch in text)
    text = unicodedata.normalize("NFKD", text)
    return text.encode("latin-1", "ignore").decode("latin-1")


def _md_to_html(markdown: str) -> str:
    """Convert a practical subset of markdown to HTML for the exporters.

    Supports headings, fenced/inline code, blockquotes, hr, single-level
    bullet/numbered lists, links, markdown tables, bold/italic, and
    paragraphs. Anything more complex should be passed as HTML content.
    """
    def inline(text: str) -> str:
        text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
        text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", text)
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', text)
        return text

    def table_block(lines: list[str]) -> str:
        head = [c.strip() for c in lines[0].strip("|").split("|")]
        rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in lines[2:]]
        html = "<table><tr>" + "".join(f"<th>{inline(h)}</th>" for h in head) + "</tr>"
        for row in rows:
            html += "<tr>" + "".join(f"<td>{inline(c)}</td>" for c in row) + "</tr>"
        return html + "</table>"

    blocks: list[str] = []
    md_lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(md_lines):
        line = md_lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            blocks.append("<hr>")
            i += 1
            continue
        if stripped.startswith("```"):
            code = []
            i += 1
            while i < len(md_lines) and not md_lines[i].strip().startswith("```"):
                code.append(md_lines[i])
                i += 1
            i += 1
            blocks.append(f"<pre>{code}</pre>")
            continue
        if stripped.startswith("#"):
            match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if match:
                level = len(match.group(1))
                blocks.append(f"<h{level}>{inline(match.group(2))}</h{level}>")
                i += 1
                continue
        if stripped.startswith(">"):
            quote = []
            while i < len(md_lines) and md_lines[i].strip().startswith(">"):
                quote.append(md_lines[i].strip().lstrip("> ").strip())
                i += 1
            blocks.append(f"<blockquote>{inline(' '.join(quote))}</blockquote>")
            continue
        if stripped.startswith("|") and i + 1 < len(md_lines) and re.match(r"^\|?[\s:|-]+-[\s:|-]+\|?$", md_lines[i + 1].strip()):
            table_lines = []
            while i < len(md_lines) and md_lines[i].strip().startswith("|"):
                table_lines.append(md_lines[i])
                i += 1
            blocks.append(table_block(table_lines))
            continue

        list_marker = re.match(r"^([-*+]|\d+\.)\s+(.*)$", stripped)
        if list_marker:
            ordered = not list_marker.group(1).startswith(("-", "*", "+"))
            open_tag, close_tag = ("<ol>", "</ol>") if ordered else ("<ul>", "</ul>")
            items = [inline(list_marker.group(2))]
            i += 1
            while i < len(md_lines):
                nested = re.match(r"^([-*+]|\d+\.)\s+(.*)$", md_lines[i].strip())
                if not nested:
                    break
                items.append(inline(nested.group(2)))
                i += 1
            blocks.append(open_tag + "".join(f"<li>{it}</li>" for it in items) + close_tag)
            continue

        paragraph = [line]
        i += 1
        while i < len(md_lines) and md_lines[i].strip() and not re.match(r"^(#{1,6})\s+", md_lines[i].strip()) and not md_lines[i].strip().startswith((">", "|", "```")):
            list_check = re.match(r"^([-*+]|\d+\.)\s+", md_lines[i].strip())
            if list_check:
                break
            paragraph.append(md_lines[i])
            i += 1
        blocks.append(f"<p>{inline(' '.join(p for p in paragraph if p.strip()))}</p>")

    return "\n".join(blocks)


def _normalize_to_html(content: str, content_type: str) -> str:
    """Return HTML regardless of input type, restricted to fpdf2-safe tags.

    Tool views (case brief, summary, ...) build their printable HTML with
    <div class="field-label"> section titles instead of heading tags, so
    those are promoted to <h3> first — otherwise PDF/DOCX exports would
    render section titles at body size instead of bold and larger.
    """
    html = content if content_type == "html" else _md_to_html(content)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all("div", class_="field-label"):
        tag.name = "h3"
    for tag in soup.find_all(["div", "section"]):
        tag.name = "p"
    for tag in soup.find_all(True):
        if tag.name not in _FPDF_SAFE_TAGS:
            tag.unwrap()
    return str(soup)


def _pdf_bytes(html: str, title: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("helvetica", "B", 15)
    pdf.multi_cell(0, 9, _sanitize_latin1(title), align="L")
    pdf.ln(3)
    pdf.set_font("helvetica", "", 11)
    pdf.write_html(_sanitize_latin1(html), table_line_separators=True)
    return bytes(pdf.output())


def _add_inline_runs(paragraph, html_text: str) -> None:
    """Add runs to a docx paragraph, honoring <b>/<i>/<u>/<code>/<a> tags."""
    wrap = BeautifulSoup(f"<root>{html_text}</root>", "html.parser").find("root")

    def walk(node, bold=False, italic=False, underline=False, code=False):
        for child in node.children:
            if getattr(child, "name", None) is None:
                text = str(child)
                if not text:
                    continue
                run = paragraph.add_run(text)
                run.bold, run.italic, run.underline = bold, italic, underline
                if code:
                    run.font.name = "Courier New"
                continue
            name = child.name
            if name in ("b", "strong"):
                walk(child, True, italic, underline, code)
            elif name in ("i", "em"):
                walk(child, bold, True, underline, code)
            elif name == "u":
                walk(child, bold, italic, True, code)
            elif name == "code":
                walk(child, bold, italic, underline, True)
            elif name == "a":
                text = child.get_text()
                if text:
                    run = paragraph.add_run(text)
                    run.bold, run.italic, run.underline = bold, italic, underline
                    run.font.color.rgb = RGBColor(0x20, 0x9D, 0xD7)
                    if code:
                        run.font.name = "Courier New"
            else:
                walk(child, bold, italic, underline, code)

    walk(wrap)


def _add_docx_element(doc: DocxDocument, el) -> None:
    """Append one HTML block element to the word document."""
    text = el.get_text()

    if el.name in ("p", "blockquote", "pre", "code"):
        paragraph = doc.add_paragraph()
        if el.name == "blockquote":
            paragraph.paragraph_format.left_indent = Pt(24)
        if el.name in ("pre", "code"):
            paragraph.style = doc.styles["Normal"]
        _add_inline_runs(paragraph, el.decode_contents())
    elif el.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        doc.add_heading(text, level=min(int(el.name[1]), 4))
    elif el.name in ("ul", "ol"):
        style = "List Number" if el.name == "ol" else "List Bullet"
        for li in el.find_all("li", recursive=False):
            _add_inline_runs(doc.add_paragraph(style=style), li.decode_contents())
    elif el.name == "hr":
        paragraph = doc.add_paragraph()
        paragraph.add_run("── ── ──").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    elif el.name == "table":
        rows = el.find_all("tr")
        if not rows:
            return
        first_row_cells = rows[0].find_all(["td", "th"])
        table = doc.add_table(rows=len(rows), cols=len(first_row_cells))
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for c_idx, cell in enumerate(cells):
                if c_idx >= len(first_row_cells):
                    break
                table.cell(r_idx, c_idx).text = cell.get_text()


def _docx_bytes(html: str, title: str) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    soup = BeautifulSoup(html, "html.parser")
    body = soup.body or soup
    for el in body.find_all(recursive=False):
        if el.name in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "blockquote", "pre", "hr"):
            _add_docx_element(doc, el)
        else:
            for child in el.find_all(recursive=False):
                if child.name in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "blockquote", "pre", "hr"):
                    _add_docx_element(doc, child)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _html_to_markdown(html: str) -> str:
    """Convert normalized HTML into editable markdown (for .md export of
    tool-view content, which only exists as HTML on the client)."""
    def inline(node, text: str = "") -> str:
        for child in node.children:
            if getattr(child, "name", None) is None:
                text += str(child)
                continue
            name = child.name
            inner = inline(child)
            if name in ("b", "strong"):
                text += f"**{inner}**"
            elif name in ("i", "em"):
                text += f"*{inner}*"
            elif name == "code":
                text += f"`{inner}`"
            elif name == "a":
                text += f"[{inner}]({child.get('href', '')})"
            elif name == "br":
                text += "\n"
            else:
                text += inner
        return text

    soup = BeautifulSoup(html, "html.parser")
    body = soup.body or soup
    lines: list[str] = []
    for el in body.find_all(recursive=False):
        if el.name in ("p", "blockquote", "pre", "code"):
            prefix = "> " if el.name == "blockquote" else ""
            lines.append(prefix + inline(el).strip())
        elif el.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            lines.append("#" * int(el.name[1]) + " " + inline(el).strip())
        elif el.name == "ul":
            for li in el.find_all("li", recursive=False):
                lines.append("- " + inline(li).strip())
        elif el.name == "ol":
            for idx, li in enumerate(el.find_all("li", recursive=False), start=1):
                lines.append(f"{idx}. " + inline(li).strip())
        elif el.name == "hr":
            lines.append("---")
        elif el.name == "table":
            rows = el.find_all("tr")
            if rows:
                head = [inline(c).strip() for c in rows[0].find_all(["td", "th"])]
                lines.append("| " + " | ".join(head) + " |")
                lines.append("| " + " | ".join("---" for _ in head) + " |")
                for row in rows[1:]:
                    cells = [inline(c).strip() for c in row.find_all(["td", "th"])]
                    lines.append("| " + " | ".join(cells) + " |")
        else:
            for child in el.find_all(recursive=False):
                if child.name in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "blockquote", "hr"):
                    lines.append(inline(child).strip())
    return "\n\n".join(line for line in lines if line)


def _prepend_title_if_missing(body: str, title: str) -> str:
    """Add the export title as an H1 unless the body already starts with one.

    PDF/DOCX always render the title as a heading; markdown exports should
    match, but content that already ships an H1 (e.g. saved documents) must
    not get a duplicated title line.
    """
    if not title or not body.strip():
        return body
    first_line = body.strip().split("\n", 1)[0]
    if first_line.startswith("# "):
        return body
    return f"# {title}\n\n{body}"


def _strip_duplicate_heading(html: str, title: str) -> str:
    """Drop the body's first heading when it repeats the export title.

    Tool views (bluebook formatter, issue spotter, case comparison) ship
    their own <h2> title inside the printable HTML, and PDF/DOCX/MD exports
    also render `filename` as a title — without this, downloads show the
    heading twice. A non-matching first heading is left untouched.
    """
    if not title:
        return html
    target = " ".join(title.split()).strip().lower()
    if not target:
        return html
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        text = " ".join(tag.get_text().split()).strip().lower()
        if text == target:
            tag.decompose()
        break
    return str(soup)


def export_content(content: str, content_type: str, fmt: str, title: str = "") -> bytes:
    """Render `content` as PDF/DOCX bytes or raw markdown bytes."""
    if fmt not in SUPPORTED_FORMATS:
        raise ExportError(f"Unsupported format '{fmt}'. Choose from {', '.join(SUPPORTED_FORMATS)}")
    if content_type not in SUPPORTED_CONTENT_TYPES:
        raise ExportError(f"Unsupported content_type '{content_type}'. Choose from {', '.join(SUPPORTED_CONTENT_TYPES)}")
    if not content or not content.strip():
        raise ExportError("Nothing to export: content is empty")

    normalized = _strip_duplicate_heading(_normalize_to_html(content, content_type), title)
    if fmt == "md":
        if content_type == "markdown":
            body = content
        else:
            body = _html_to_markdown(normalized)
        return _prepend_title_if_missing(body, title).encode("utf-8")
    if fmt == "pdf":
        return _pdf_bytes(normalized, title or "Document")
    return _docx_bytes(normalized, title or "Document")


def sanitize_filename(name: str) -> str:
    """Strip path separators, dots, and control characters from a user filename."""
    cleaned = re.sub(r"[^a-zA-Z0-9_\- ]", "", name)
    cleaned = re.sub(r"\s+", "_", cleaned).strip("._- ") or "document"
    return cleaned