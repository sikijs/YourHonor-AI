import os
from typing import Optional
from pathlib import Path
from langchain_core.documents import Document

from .chunking import LegalTextSplitter
from .qdrant_store import add_documents


class IngestionService:
    def __init__(self):
        self.splitter = LegalTextSplitter()

    def load_text_file(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def load_markdown_file(self, file_path: str) -> str:
        return self.load_text_file(file_path)

    def load_pdf(self, file_path: str) -> Optional[str]:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages])
            return text
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return None

    def ocr_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from a scanned (image-only) PDF via tesseract OCR.

        Each page is rendered to an image with PyMuPDF, then run through
        pytesseract. Slower than load_pdf and sensitive to scan quality, so
        it is used only as a fallback when a PDF has no embedded text.
        """
        try:
            import io
            import pymupdf
            import pytesseract
            from PIL import Image

            pages_text = []
            with pymupdf.open(file_path) as doc:
                for page in doc:
                    pix = page.get_pixmap(dpi=200)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    pages_text.append(pytesseract.image_to_string(img))
            text = "\n".join(pages_text)
            return text if text.strip() else None
        except Exception as e:
            print(f"Error running OCR on PDF {file_path}: {e}")
            return None

    def load_docx(self, file_path: str) -> Optional[str]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return None

    def ingest_document(
        self,
        content: str,
        title: str,
        source: str = "user_upload",
        metadata: Optional[dict] = None,
    ) -> dict:
        doc = Document(
            page_content=content,
            metadata={
                "title": title,
                "source": source,
                **(metadata or {}),
            },
        )
        
        chunks = self.splitter.split_documents([doc])
        
        add_documents(chunks)
        
        return {
            "title": title,
            "source": source,
            "chunks_created": len(chunks),
        }

    def ingest_file(
        self,
        file_path: str,
        title: Optional[str] = None,
        metadata: Optional[dict] = None,
    ) -> dict:
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == ".txt":
            content = self.load_text_file(file_path)
        elif extension == ".md":
            content = self.load_markdown_file(file_path)
        elif extension == ".pdf":
            content = self.load_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            content = self.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        if content is None:
            raise ValueError(f"Failed to extract content from {file_path}")
        
        doc_title = title or path.stem
        return self.ingest_document(content, doc_title, str(path), metadata)

    def ingest_directory(
        self,
        directory_path: str,
        file_extensions: list[str] = [".md", ".txt"],
    ) -> list[dict]:
        path = Path(directory_path)
        results = []
        
        for ext in file_extensions:
            for file_path in path.glob(f"*{ext}"):
                try:
                    result = self.ingest_file(str(file_path))
                    results.append(result)
                except Exception as e:
                    print(f"Error ingesting {file_path}: {e}")
        
        return results


_ingestion_service: Optional[IngestionService] = None


def get_ingestion_service() -> IngestionService:
    global _ingestion_service
    if _ingestion_service is None:
        _ingestion_service = IngestionService()
    return _ingestion_service